import datetime
from docx import Document
from docx.shared import  RGBColor
from docx.oxml.ns import qn
# 新增：生成 Word 文档
def create_word_document(chapters, summary_data, meetingassistance_data, transcription_data):
    document = Document()

    def set_font(run, font_name='等线', font_size=None, bold=False, italic=False, color=RGBColor(0, 0, 0)):
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.color.rgb = color
        run.font.bold = bold
        run.font.italic = italic
        if font_size:
            run.font.size = font_size

    def add_heading(text, level):
        heading = document.add_heading(text, level=level)
        run = heading.runs[0]
        set_font(run, bold=True)

    def add_paragraph(text):
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text)  # 确保有一个 run 可以设置字体
        set_font(run)

    add_heading('会议纪要', 0)

    add_heading('关键词', level=1)
    add_paragraph(meetingassistance_data['keywords'])

    add_heading('全文摘要', level=1)
    if "paragraph_summary" in summary_data:
        add_paragraph(summary_data['paragraph_summary'])

    add_heading('章节速览', level=1)
    for chapter in chapters:
        add_heading(chapter['title'], level=2)
        add_paragraph(chapter['summary'])

    add_heading('发言总结', level=1)
    for conv in summary_data['conversational_summary']:
        add_heading(f"发言人 {conv['speaker_id']} - {conv['speaker_name']}", level=2)
        add_paragraph(conv['summary'])

    add_heading('问答总结', level=1)
    for qa in summary_data['qa_summary']:
        add_heading(f"问题: {qa['question']}", level=2)
        add_paragraph(f"答案: {qa['answer']}")

    add_heading('要点回顾', level=1)
    add_paragraph(meetingassistance_data['key_sentences'])

    add_heading('待办事项', level=1)
    add_paragraph(meetingassistance_data['actions'])

    add_heading('全文记录', level=1)
    add_paragraph(transcription_data['transcription'])

    # 使用当前时间作为文件名
    current_time = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    save_path = f"output/会议纪要_{current_time}.docx"
    document.save(save_path)
    return save_path